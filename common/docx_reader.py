from docx import Document

def extract_text_from_docx(path):
    doc = Document(path)
    lines = []

    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())

    for t in doc.tables:
        for row in t.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)

    return "\n".join(lines)